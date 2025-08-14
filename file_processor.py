
import streamlit as st
import pandas as pd
import io
from typing import Union
import mimetypes

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        USE_PDFPLUMBER = True
    except ImportError:
        PDF_AVAILABLE = False
        USE_PDFPLUMBER = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class FileProcessor:
    """Handles processing of different file types and extracting text content."""
    
    def __init__(self):
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit
        
    def process_file(self, uploaded_file) -> str:
        """
        Process uploaded file and extract text content.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content from the file
            
        Raises:
            Exception: If file processing fails
        """
        # Check file size
        if uploaded_file.size > self.max_file_size:
            raise Exception(f"File size ({uploaded_file.size} bytes) exceeds maximum limit ({self.max_file_size} bytes)")
        
        # Get file extension
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        # Process based on file type
        try:
            if file_extension == 'txt':
                return self._process_text_file(uploaded_file)
            elif file_extension == 'pdf':
                return self._process_pdf_file(uploaded_file)
            elif file_extension == 'docx':
                return self._process_docx_file(uploaded_file)
            elif file_extension in ['xlsx', 'xls']:
                return self._process_excel_file(uploaded_file)
            elif file_extension == 'csv':
                return self._process_csv_file(uploaded_file)
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Error processing {file_extension.upper()} file: {str(e)}")
    
    def _process_text_file(self, uploaded_file) -> str:
        """Process plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    content = uploaded_file.read().decode(encoding)
                    uploaded_file.seek(0)  # Reset file pointer
                    return content
                except UnicodeDecodeError:
                    uploaded_file.seek(0)  # Reset file pointer
                    continue
            
            raise Exception("Unable to decode file with common encodings")
            
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    def _process_pdf_file(self, uploaded_file) -> str:
        """Process PDF files."""
        if not PDF_AVAILABLE:
            raise Exception("PDF processing libraries not available. Please install PyPDF2 or pdfplumber.")
        
        try:
            content = ""
            
            if 'USE_PDFPLUMBER' in globals() and USE_PDFPLUMBER:
                # Use pdfplumber
                import pdfplumber
                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
            else:
                # Use PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"
            
            if not content.strip():
                raise Exception("No text content found in PDF")
            
            return content
            
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    def _process_docx_file(self, uploaded_file) -> str:
        """Process Word documents."""
        if not DOCX_AVAILABLE:
            raise Exception("DOCX processing library not available. Please install python-docx.")
        
        try:
            doc = Document(uploaded_file)
            content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            if not content.strip():
                raise Exception("No text content found in DOCX file")
            
            return content
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _process_excel_file(self, uploaded_file) -> str:
        """Process Excel files."""
        if not EXCEL_AVAILABLE:
            raise Exception("Excel processing library not available. Please install openpyxl.")
        
        try:
            # Read Excel file
            df_dict = pd.read_excel(uploaded_file, sheet_name=None)
            content = ""
            
            for sheet_name, df in df_dict.items():
                content += f"\n--- Sheet: {sheet_name} ---\n"
                content += df.to_string(index=False) + "\n"
            
            if not content.strip():
                raise Exception("No data found in Excel file")
            
            return content
            
        except Exception as e:
            raise Exception(f"Error reading Excel file: {str(e)}")
    
    def _process_csv_file(self, uploaded_file) -> str:
        """Process CSV files."""
        try:
            # Try different encodings and separators
            encodings = ['utf-8', 'latin-1', 'cp1252']
            separators = [',', ';', '\t']
            
            for encoding in encodings:
                for sep in separators:
                    try:
                        uploaded_file.seek(0)
                        df = pd.read_csv(uploaded_file, encoding=encoding, sep=sep)
                        
                        # Check if DataFrame has reasonable structure
                        if len(df.columns) > 1 or len(df) > 0:
                            content = df.to_string(index=False)
                            if content.strip():
                                return content
                                
                    except Exception:
                        continue
            
            # If all else fails, try reading as plain text
            uploaded_file.seek(0)
            return self._process_text_file(uploaded_file)
            
        except Exception as e:
            raise Exception(f"Error reading CSV file: {str(e)}")
    
    def get_supported_types(self) -> list:
        """Return list of supported file types."""
        supported = ['txt', 'csv']
        
        if PDF_AVAILABLE:
            supported.append('pdf')
        if DOCX_AVAILABLE:
            supported.append('docx')
        if EXCEL_AVAILABLE:
            supported.extend(['xlsx', 'xls'])
            
        return supported
    
    def get_missing_dependencies(self) -> list:
        """Return list of missing dependencies."""
        missing = []
        
        if not PDF_AVAILABLE:
            missing.append('PyPDF2 or pdfplumber (for PDF files)')
        if not DOCX_AVAILABLE:
            missing.append('python-docx (for Word documents)')
        if not EXCEL_AVAILABLE:
            missing.append('openpyxl (for Excel files)')
            
        return missing
